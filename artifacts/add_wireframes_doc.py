# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path


ROOT = Path(r"D:\developSoftware\ideacode\ai-dinner-frontend-main")
SRC_DIR = Path(r"D:\study\ai-dinner")
src_docs = sorted(SRC_DIR.glob("*.docx"))
if not src_docs:
    raise FileNotFoundError(f"No .docx found in {SRC_DIR}")

SRC_DOC = src_docs[0]
OUT_DIR = ROOT / "artifacts" / "wireframes"
OUT_DOC = ROOT / "artifacts" / "ai-dinner-design-with-wireframes.docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

WIDTH, HEIGHT = 390, 844
BG = "#F7EFE5"
CARD = "#FFF9F1"
CARD_ALT = "#FFFFFF"
INK = "#2D2925"
MUTED = "#7D746A"
LINE = "#D8C9B8"
ACCENT = "#A85F2D"
ACCENT_SOFT = "#F0D6C0"
GREEN = "#5F7B5A"
BLUE = "#5C7284"
RED = "#A46363"

font_candidates = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
    Path(r"C:\Windows\Fonts\arial.ttf"),
]
FONT_PATH = next((p for p in font_candidates if p.exists()), None)


def font(size):
    return ImageFont.truetype(str(FONT_PATH), size) if FONT_PATH else ImageFont.load_default()


F_TITLE = font(24)
F_H2 = font(18)
F_BODY = font(13)
F_SMALL = font(11)
F_TAG = font(12)
F_BTN = font(14)


def rounded(draw, xy, fill, outline=LINE, radius=16, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_text(draw, xy, value, fill=INK, fnt=F_BODY, max_width=None, line_gap=4):
    x, y = xy
    if not max_width:
        draw.text((x, y), value, fill=fill, font=fnt)
        return y + 16

    lines = []
    cur = ""
    for ch in value:
        test = cur + ch
        if draw.textlength(test, font=fnt) <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = ch
    if cur:
        lines.append(cur)

    for line in lines:
        draw.text((x, y), line, fill=fill, font=fnt)
        box = draw.textbbox((x, y), line, font=fnt)
        y += (box[3] - box[1]) + line_gap
    return y


def chip(draw, x, y, label, selected=False, color=None):
    w = int(draw.textlength(label, font=F_TAG)) + 20
    h = 28
    fill = color or (ACCENT_SOFT if selected else "#FFFFFF")
    outline = ACCENT if selected else LINE
    rounded(draw, (x, y, x + w, y + h), fill=fill, outline=outline, radius=14)
    draw.text((x + 10, y + 6), label, fill=ACCENT if selected else MUTED, font=F_TAG)
    return x + w + 8


def section_label(draw, y, label):
    draw.text((22, y), label, fill=MUTED, font=F_SMALL)
    return y + 20


def phone_base(title):
    img = Image.new("RGB", (WIDTH, HEIGHT), BG)
    draw = ImageDraw.Draw(img)
    draw.text((22, 26), title, fill=INK, font=F_TITLE)
    return img, draw


def save_image(img, name):
    path = OUT_DIR / name
    img.save(path)
    return path


def draw_home():
    img, draw = phone_base("今晚想吃点什么？")
    rounded(draw, (22, 76, 368, 138), CARD)
    draw_text(draw, (38, 92), "我记得你最近偏爱微辣、热乎、20 分钟内完成。", MUTED, F_BODY, 300)
    rounded(draw, (22, 154, 368, 256), CARD_ALT)
    draw_text(draw, (38, 172), "有鸡蛋、番茄、豆腐，今天很累，想吃热乎一点，最好少洗锅。", "#4D4740", F_BODY, 298)
    draw.line((38, 232, 160, 232), fill=LINE, width=1)
    y = 278
    for label, tags in [
        ("今天状态", ["累了", "没胃口", "想奖励自己"]),
        ("口味偏好", ["微辣", "热乎", "下饭", "少洗锅"]),
        ("做饭时间", ["10 分钟", "20 分钟", "30 分钟"]),
        ("可用厨具", ["一口炒锅", "电饭锅", "空气炸锅"]),
    ]:
        y = section_label(draw, y, label)
        x = 22
        for i, tag in enumerate(tags):
            x = chip(draw, x, y, tag, selected=i in (0, 1))
        y += 44
    rounded(draw, (22, 742, 368, 794), ACCENT, outline=ACCENT, radius=18)
    draw.text((105, 759), "帮我想想今晚吃啥", fill="white", font=F_BTN)
    return save_image(img, "01-home.png")


def reco_card(draw, y, kind, name, reason, color):
    rounded(draw, (22, y, 368, y + 185), CARD_ALT)
    rounded(draw, (38, y + 18, 98, y + 44), color, outline=color, radius=13)
    draw.text((49, y + 24), kind, fill="white", font=F_SMALL)
    draw.text((110, y + 18), name, fill=INK, font=F_H2)
    draw_text(draw, (38, y + 54), reason, MUTED, F_SMALL, 294, 3)
    draw.text((38, y + 94), "18 分钟 · 简单 · 用到：鸡蛋/番茄/米饭", fill="#59514A", font=F_SMALL)
    draw.text((38, y + 118), "缺少：葱花（可选）", fill=RED, font=F_SMALL)
    draw_text(draw, (38, y + 140), "步骤：炒蛋 → 炒番茄 → 合锅调味 → 盖饭", "#4D4740", F_SMALL, 290, 2)
    chip(draw, 38, y + 156, "想吃", selected=True, color="#E9F0E7")
    chip(draw, 98, y + 156, "收藏")
    chip(draw, 158, y + 156, "太麻烦")


def draw_result():
    img, draw = phone_base("今晚可以这样吃")
    draw_text(draw, (22, 60), "根据“很累、想热乎、少洗锅”给你 3 个方向。", MUTED, F_BODY, 330)
    reco_card(draw, 96, "最省事", "番茄鸡蛋热汤面", "不用额外买菜，一口锅就能完成，适合下班后快速吃上热乎的。", ACCENT)
    reco_card(draw, 302, "最满足", "微辣番茄鸡蛋盖饭", "更下饭，饱腹感强，能照顾今天想吃有味道一点的状态。", BLUE)
    reco_card(draw, 508, "最健康", "豆腐蛋花番茄汤", "少油清爽，有蛋白质和热汤，适合不想吃太重的时候。", GREEN)
    return save_image(img, "02-result.png")


def pref_row(draw, y, label, value):
    draw.text((38, y), label, fill=MUTED, font=F_SMALL)
    draw_text(draw, (132, y - 2), value, INK, F_BODY, 195, 3)


def draw_profile():
    img, draw = phone_base("我的口味")
    rounded(draw, (22, 78, 368, 164), CARD)
    draw_text(draw, (38, 96), "你最近更喜欢：微辣、热乎、下饭、20 分钟内完成。", INK, F_BODY, 292)
    draw.text((38, 136), "画像会根据反馈持续更新，也可以手动调整。", fill=MUTED, font=F_SMALL)
    y = 188
    rounded(draw, (22, y, 368, y + 310), CARD_ALT)
    pref_row(draw, y + 24, "常用食材", "鸡蛋、番茄、豆腐、青菜、剩米饭")
    pref_row(draw, y + 74, "不喜欢", "香菜、肥肉、内脏")
    pref_row(draw, y + 124, "做饭偏好", "少洗锅、一锅出、简单调味、最多 3 步")
    pref_row(draw, y + 184, "辣度", "微辣，可接受偶尔中辣")
    pref_row(draw, y + 226, "健康目标", "晚餐少油、蛋白质够、不要太撑")
    y = 526
    draw.text((22, y), "快捷调整", fill=MUTED, font=F_SMALL)
    y += 24
    for row in [["少油", "高蛋白", "低碳水"], ["少洗锅", "不吃香菜", "空气炸锅"]]:
        x = 22
        for tag in row:
            x = chip(draw, x, y, tag, selected=tag in ["少油", "少洗锅", "不吃香菜"])
        y += 44
    rounded(draw, (22, 742, 368, 794), ACCENT, outline=ACCENT, radius=18)
    draw.text((159, 759), "保存", fill="white", font=F_BTN)
    return save_image(img, "03-profile.png")


def fav_card(draw, y, name, tags):
    rounded(draw, (22, y, 368, y + 126), CARD_ALT)
    draw.text((38, y + 18), name, fill=INK, font=F_H2)
    draw.text((38, y + 48), "上次反馈：想吃 · 今天做了", fill=MUTED, font=F_SMALL)
    x = 38
    for tag in tags:
        x = chip(draw, x, y + 70, tag)
    rounded(draw, (38, y + 102, 158, y + 122), "#FFFFFF", outline=ACCENT, radius=10)
    draw.text((58, y + 106), "再来类似的", fill=ACCENT, font=F_SMALL)
    rounded(draw, (176, y + 102, 326, y + 122), ACCENT_SOFT, outline=ACCENT_SOFT, radius=10)
    draw.text((210, y + 106), "加入今日晚餐", fill=ACCENT, font=F_SMALL)


def draw_favorites():
    img, draw = phone_base("收藏")
    x = 22
    for tag in ["全部", "快手", "下饭", "少油", "热乎"]:
        x = chip(draw, x, 74, tag, selected=tag == "全部")
    fav_card(draw, 126, "番茄鸡蛋热汤面", ["热乎", "快手", "一口锅"])
    fav_card(draw, 270, "空气炸锅鸡腿饭", ["满足", "空气炸锅"])
    fav_card(draw, 414, "豆腐蛋花番茄汤", ["少油", "清淡", "高蛋白"])
    fav_card(draw, 558, "青菜火腿炒饭", ["下饭", "剩饭友好"])
    return save_image(img, "04-favorites.png")


def hist_group(draw, y, date, input_text, reco, feedback):
    draw.text((22, y), date, fill=MUTED, font=F_SMALL)
    rounded(draw, (22, y + 22, 368, y + 154), CARD_ALT)
    draw_text(draw, (38, y + 40), "输入：" + input_text, "#4D4740", F_SMALL, 292, 3)
    draw_text(draw, (38, y + 76), "推荐：" + reco, INK, F_BODY, 292, 3)
    draw.text((38, y + 110), "反馈：" + feedback, fill=MUTED, font=F_SMALL)
    rounded(draw, (236, y + 112, 344, y + 138), ACCENT_SOFT, outline=ACCENT_SOFT, radius=13)
    draw.text((258, y + 119), "再来类似的", fill=ACCENT, font=F_SMALL)


def draw_history():
    img, draw = phone_base("历史")
    hist_group(draw, 82, "今天", "有鸡蛋番茄，今天很累，想热乎。", "番茄鸡蛋热汤面、微辣盖饭、豆腐蛋花汤", "收藏 · 今天做了")
    hist_group(draw, 258, "昨天", "想吃清淡，只有青菜和豆腐。", "青菜豆腐汤、豆腐煎蛋、清炒青菜盖饭", "太清淡")
    hist_group(draw, 434, "5 月 14 日", "想奖励自己，30 分钟内。", "空气炸锅鸡腿饭、酸辣汤面、照烧豆腐饭", "想吃")
    hist_group(draw, 610, "5 月 13 日", "没胃口，想酸甜一点。", "番茄肥牛汤、酸甜鸡蛋饭、凉拌黄瓜豆腐", "不想吃")
    return save_image(img, "05-history.png")


def main():
    image_paths = [draw_home(), draw_result(), draw_profile(), draw_favorites(), draw_history()]
    doc = Document(str(SRC_DOC))
    body = doc._body._element
    insert_before = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("6. 前端请求封装"):
            insert_before = para._p
            break
    if insert_before is None:
        insert_before = body[-1]

    def move_before(element):
        body.insert(body.index(insert_before), element)

    def add_para(value="", style=None, size=None, color=None, align=None):
        para = doc.add_paragraph(style=style)
        if value:
            run = para.add_run(value)
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
        if align is not None:
            para.alignment = align
        move_before(para._p)
        return para

    def add_bullet(value):
        add_para(value, style="List Bullet")

    def add_picture(path, caption):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(path), width=Inches(2.75))
        move_before(para._p)
        cap = add_para(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(125, 116, 106)

    def add_table(rows):
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, value in enumerate(rows[0]):
            table.rows[0].cells[i].text = value
        for row in rows[1:]:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        move_before(table._tbl)

    add_para("5.7 低保真移动端页面设计（iPhone 390px）", style="Heading 2")
    add_para("本节用于指导 uni-app + Vue3 前端实现首版微信小程序页面。低保真稿按 iPhone 390px 宽度设计，强调决策助手体验，而不是传统菜谱检索。页面应让用户感觉是在和一个懂自己、能帮忙做决定的 AI 饭搭子对话。")
    add_para("Figma 协作文件： https://www.figma.com/design/vPc9ewEBl44HbguoIUj0I5 。当前文档中的线框图和组件说明为前端实现依据。")
    add_para("设计原则", style="Heading 3")
    for item in [
        "画布尺寸：390px 宽，建议首屏高度按 844px 设计；小程序实际实现使用 rpx/vw 适配，核心内容宽度保持 342px 左右。",
        "页面边距：左右 22px；卡片内边距 16px；模块间距 16px；同组元素间距 8px 到 12px。",
        "视觉风格：暖米色背景、白色或浅暖色卡片、低饱和橙棕作为主按钮和选中态，整体轻松生活化。",
        "组件风格：卡片圆角 16px；标签 pill 圆角 14px；主按钮高度 52px；尽量用自然语言解释推荐原因。",
        "信息层级：每页只保留一个核心任务；弱化菜谱感，突出“今天状态 + 食材 + 决策理由 + 反馈学习”。",
    ]:
        add_bullet(item)

    add_para("基础视觉 Token", style="Heading 3")
    add_table([
        ["角色", "建议值", "说明"],
        ["pageBg", "#F7EFE5", "晚餐场景的暖米色背景"],
        ["cardBg", "#FFF9F1 / #FFFFFF", "提示卡片和内容卡片"],
        ["primary", "#A85F2D", "主按钮、选中态、重要操作"],
        ["primarySoft", "#F0D6C0", "弱按钮和标签选中背景"],
        ["textPrimary", "#2D2925", "标题和主要正文"],
        ["textSecondary", "#7D746A", "说明、日期、辅助信息"],
        ["border", "#D8C9B8", "卡片描边和分割线"],
        ["radiusCard", "16px", "主要内容卡片"],
        ["radiusPill", "14px", "标签、筛选项、轻按钮"],
    ])

    pages = [
        ("页面 1：首页", "pages/index/index.vue", image_paths[0], "低保真线框图 1：首页（390px 宽）", "让用户用最少输入表达“家里有什么、今天状态、口味、时间和厨具”，并发起一次晚餐决策。", ["顶部标题：今晚想吃点什么？", "个性化提示卡片：根据新老用户展示不同语气提示。", "自然语言输入框：支持长句输入，承接真实场景表达。", "标签区：心情、口味、时间、厨具四组标签。", "底部主按钮：帮我想想今晚吃啥。"], "主按钮只有在自然语言输入或至少一个标签被选择后进入可点击态；点击后进入加载态，再跳转推荐结果页。"),
        ("页面 2：推荐结果页", "pages/result/result.vue", image_paths[1], "低保真线框图 2：推荐结果页（390px 宽）", "用 3 个明确方案降低选择成本，分别覆盖最省事、最满足、最健康。", ["顶部标题：今晚可以这样吃。", "推荐解释：一句话复述本次输入的状态和偏好。", "3 张推荐卡片：最省事、最满足、最健康。", "卡片字段：菜名、推荐理由、预计时间、难度、已用食材、缺少食材、步骤、替换建议、标签。", "反馈按钮：想吃、不想吃、太麻烦、太清淡、太油、收藏、今天做了。"], "每张卡片先给结论和理由，再给步骤。反馈按钮应即时更新本地状态，并异步提交后端用于画像学习。"),
        ("页面 3：我的口味页", "pages/profile/profile.vue", image_paths[2], "低保真线框图 3：我的口味页（390px 宽）", "让用户看到 AI 正在学习自己，也允许用户修正画像。", ["口味画像摘要：用一句话描述最近偏好。", "常用食材：展示高频使用食材。", "不喜欢食材：展示明确排除项。", "做饭偏好：时间、难度、厨具、少洗锅、一锅出等。", "辣度和健康目标：用于影响后续推荐。", "保存按钮：提交用户手动调整。"], "画像摘要应使用可解释文案，不直接暴露算法分数。保存后展示轻提示，并刷新 profile store。"),
        ("页面 4：收藏页", "pages/favorites/favorites.vue", image_paths[3], "低保真线框图 4：收藏页（390px 宽）", "让用户快速复用喜欢的菜，并从收藏中生成今天的晚餐方向。", ["标签筛选：全部、快手、下饭、少油、热乎等。", "收藏卡片：菜名、上次反馈、标签。", "操作按钮：再来类似的、加入今日晚餐。", "空状态：提示先去推荐页收藏一个喜欢的方案。"], "“再来类似的”带 favoriteId 请求推荐接口；“加入今日晚餐”可进入确认态或直接记录为今天做了。"),
        ("页面 5：历史页", "pages/history/history.vue", image_paths[4], "低保真线框图 5：历史页（390px 宽）", "按日期回看过去输入、推荐和反馈，让用户感知系统越来越懂自己。", ["日期分组：今天、昨天、具体日期。", "输入摘要：保留用户当时的食材和状态。", "推荐结果摘要：展示当次 3 个方案名称。", "反馈摘要：收藏、今天做了、不想吃、太麻烦等。", "操作按钮：再来类似的。"], "历史页不需要展示完整步骤，避免信息过重；点击卡片可进入详情，点击按钮直接发起相似推荐。"),
    ]
    for title, route, image_path, caption, goal, items, note in pages:
        add_para(title, style="Heading 3")
        add_picture(image_path, caption)
        add_para("路由：" + route)
        add_para("页面目标：" + goal)
        add_para("结构拆解：")
        for item in items:
            add_bullet(item)
        add_para("交互与实现要点：" + note)

    add_para("前端实现约束", style="Heading 3")
    for item in [
        "页面组件优先复用：TagSelector、RecommendationCard、PreferenceEditor、FavoriteCard、HistoryGroup。",
        "首页标签状态统一维护为数组或单选值，生成请求保持与 GenerateDinnerRequest 一致。",
        "推荐结果卡片应支持 loading、error、empty、feedbackSubmitted 四类状态。",
        "画像、收藏、历史页面优先从 Pinia store 读取缓存，再发起接口刷新，避免小程序页面切换时闪烁。",
        "所有“再来类似的”入口最终都转化为一次带上下文的 recommendations/generate 请求。",
    ]:
        add_bullet(item)

    doc.save(str(OUT_DOC))
    print(OUT_DOC)


if __name__ == "__main__":
    main()
